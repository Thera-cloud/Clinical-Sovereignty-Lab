"""
Sovereign Vault — File upload pipeline (B2).

FileProcessor: MIME validation, text extraction, image processing, thumbnails.
"""

import logging
import mimetypes
import hashlib
import base64
import io
import os
from dataclasses import dataclass, field
from typing import Optional, List

try:
    import magic
except (ImportError, OSError):
    magic = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PIL import Image
except ImportError:
    Image = None

SUPPORTED_TYPES = {
    "application/pdf": {"parser": "pymupdf", "max_size_mb": 25, "extract": "text"},
    "text/plain": {"parser": "direct", "max_size_mb": 10, "extract": "text"},
    "text/markdown": {"parser": "direct", "max_size_mb": 10, "extract": "text"},
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {
        "parser": "docx",
        "max_size_mb": 25,
        "extract": "text",
    },
    "image/jpeg": {"parser": "vision", "max_size_mb": 20, "extract": "base64"},
    "image/png": {"parser": "vision", "max_size_mb": 20, "extract": "base64"},
    "image/webp": {"parser": "vision", "max_size_mb": 20, "extract": "base64"},
    "image/gif": {"parser": "vision", "max_size_mb": 10, "extract": "base64"},
    "application/zip": {"parser": "transfer", "max_size_mb": 200, "extract": "json"},
    "application/json": {"parser": "transfer", "max_size_mb": 100, "extract": "json"},
}

BLOCKED_TYPES = [
    "application/x-executable",
    "application/x-msdownload",
    "application/x-sh",
    "application/javascript",
    "application/vnd.ms-excel.sheet.macroEnabled.12",
    "application/vnd.ms-word.document.macroEnabled.12",
]

MAX_EXTRACTED_CHARS = 50_000
MAX_IMAGE_DIMENSION = 4096

# Fallback MIME mapping when libmagic fails (e.g., Windows, Docker without libmagic1)
_EXTENSION_MIME = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".zip": "application/zip",
    ".json": "application/json",
}

# Tier storage limits (bytes)
TIER_LIMITS_BYTES = {
    "threshold": 1 * 1024 * 1024 * 1024,  # 1 GB
    "inner_chamber": 10 * 1024 * 1024 * 1024,  # 10 GB
    "sovereign_circle": 50 * 1024 * 1024 * 1024,  # 50 GB
}
DEFAULT_TIER_LIMIT = TIER_LIMITS_BYTES["inner_chamber"]


@dataclass
class ProcessedFile:
    """Result of file processing pipeline."""

    type: str  # "document" or "image"
    text: Optional[str] = None
    base64_data: Optional[str] = None
    media_type: Optional[str] = None
    dimensions: Optional[dict] = None
    page_count: Optional[int] = None
    preview: Optional[str] = None
    thumbnail_bytes: Optional[bytes] = None
    hash: Optional[str] = None
    filename: Optional[str] = None
    flags: List[str] = field(default_factory=list)
    raw_bytes: Optional[bytes] = None
    size_bytes: int = 0


class FileProcessor:
    """
    Processes uploaded files: validates MIME, extracts content, generates thumbnails.
    """

    def validate_mime(self, file_bytes: bytes, filename: Optional[str] = None) -> str:
        """
        Detect MIME type from magic bytes using python-magic.
        Falls back to mimetypes/extension when libmagic fails (e.g., Windows, Docker).
        Raises ValueError if blocked or unsupported.
        """
        if not file_bytes:
            raise ValueError("Empty file")
        mime = None
        try:
            mime = magic.from_buffer(file_bytes, mime=True)
        except Exception as e:
            logging.getLogger(__name__).warning(
                "libmagic.from_buffer failed (%s); using mimetypes/extension fallback. "
                "Install libmagic1 for accurate detection.",
                e,
            )
        if mime and isinstance(mime, str):
            mime = mime.strip().lower()
        if not mime and filename:
            guessed, _ = mimetypes.guess_type(filename)
            if guessed:
                mime = guessed.strip().lower()
        if not mime and filename:
            ext = os.path.splitext(filename)[1].lower()
            mime = _EXTENSION_MIME.get(ext)
        if not mime or not isinstance(mime, str):
            raise ValueError("Could not detect MIME type (libmagic failed and no filename fallback)")
        if mime in BLOCKED_TYPES:
            raise ValueError(f"File type blocked for security: {mime}")
        if mime not in SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {mime}")
        return mime

    def validate_size(
        self, size_bytes: int, mime_type: str, tier: str, current_usage_bytes: int = 0
    ) -> None:
        """
        Check file size against per-type and tier limits.
        Raises ValueError if over limit.
        """
        if size_bytes <= 0:
            raise ValueError("Invalid file size")
        spec = SUPPORTED_TYPES.get(mime_type)
        if spec:
            max_mb = spec.get("max_size_mb", 10)
            max_bytes = max_mb * 1024 * 1024
            if size_bytes > max_bytes:
                raise ValueError(
                    f"File exceeds {max_mb} MB limit for {mime_type}"
                )
        tier_limit = TIER_LIMITS_BYTES.get(
            tier.lower() if isinstance(tier, str) else "inner_chamber",
            DEFAULT_TIER_LIMIT,
        )
        if current_usage_bytes + size_bytes > tier_limit:
            raise ValueError(
                f"Storage limit exceeded for tier {tier}. "
                f"Current: {current_usage_bytes}, adding: {size_bytes}, limit: {tier_limit}"
            )

    def process(
        self, filename: str, file_bytes: bytes, mime_type: str
    ) -> ProcessedFile:
        """
        Main pipeline: validate, route to extractor, compute hash.
        For transfer types (zip/json) returns minimal ProcessedFile with raw_bytes.
        """
        detected = self.validate_mime(file_bytes, filename)
        if mime_type and mime_type.lower() != detected.lower():
            mime_type = detected
        spec = SUPPORTED_TYPES.get(detected)
        if not spec:
            raise ValueError(f"Unsupported type: {detected}")
        file_hash = self.compute_hash(file_bytes)
        parser = spec.get("parser", "direct")
        extract = spec.get("extract", "text")
        pf: ProcessedFile
        if parser == "pymupdf":
            pf = self.extract_pdf(file_bytes)
        elif parser == "docx":
            pf = self.extract_docx(file_bytes)
        elif parser == "direct":
            pf = self.extract_text(file_bytes)
        elif parser == "vision":
            pf = self.process_image(file_bytes)
        else:
            # transfer: zip, json — store as-is
            pf = ProcessedFile(
                type="document",
                raw_bytes=file_bytes,
                size_bytes=len(file_bytes),
                media_type=detected,
            )
        pf.hash = file_hash
        pf.filename = filename
        pf.size_bytes = len(file_bytes)
        pf.media_type = detected
        if pf.thumbnail_bytes is None and parser in ("pymupdf", "vision"):
            try:
                pf.thumbnail_bytes = self.generate_thumbnail(
                    file_bytes, media_type=detected
                )
            except (OSError, ValueError, TypeError) as e:
                logging.getLogger(__name__).debug(
                    "Thumbnail generation skipped: %s", e
                )
        return pf

    def extract_pdf(self, file_bytes: bytes) -> ProcessedFile:
        """Extract text from PDF via PyMuPDF, truncate to MAX_EXTRACTED_CHARS."""
        text_parts: List[str] = []
        page_count = 0
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            page_count = len(doc)
            for i in range(page_count):
                page = doc.load_page(i)
                text_parts.append(page.get_text())
            doc.close()
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {e}") from e
        full = "\n\n".join(text_parts)
        if len(full) > MAX_EXTRACTED_CHARS:
            full = full[:MAX_EXTRACTED_CHARS] + "\n[...truncated]"
            flags = ["truncated"]
        else:
            flags = []
        preview = full[:500].strip() if full else None
        return ProcessedFile(
            type="document",
            text=full or None,
            preview=preview,
            page_count=page_count,
            media_type="application/pdf",
            flags=flags,
        )

    def extract_docx(self, file_bytes: bytes) -> ProcessedFile:
        """Extract text from DOCX via python-docx."""
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            paras = [p.text for p in doc.paragraphs]
            full = "\n\n".join(paras)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text for c in row.cells]
                    full += "\n" + " | ".join(cells)
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {e}") from e
        if len(full) > MAX_EXTRACTED_CHARS:
            full = full[:MAX_EXTRACTED_CHARS] + "\n[...truncated]"
            flags = ["truncated"]
        else:
            flags = []
        preview = full[:500].strip() if full else None
        return ProcessedFile(
            type="document",
            text=full or None,
            preview=preview,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            flags=flags,
        )

    def extract_text(self, file_bytes: bytes) -> ProcessedFile:
        """Direct text read (plain, markdown)."""
        try:
            raw = file_bytes.decode("utf-8", errors="replace")
        except Exception as e:
            raise ValueError(f"Text decode failed: {e}") from e
        if len(raw) > MAX_EXTRACTED_CHARS:
            raw = raw[:MAX_EXTRACTED_CHARS] + "\n[...truncated]"
            flags = ["truncated"]
        else:
            flags = []
        preview = raw[:500].strip() if raw else None
        return ProcessedFile(
            type="document",
            text=raw or None,
            preview=preview,
            flags=flags,
        )

    def process_image(self, file_bytes: bytes) -> ProcessedFile:
        """
        EXIF strip, resize to max 2048px, base64 encode, dimension check 4096 max.
        """
        try:
            img = Image.open(io.BytesIO(file_bytes))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            elif img.mode != "RGB":
                img = img.convert("RGB")
        except Exception as e:
            raise ValueError(f"Image open failed: {e}") from e
        w, h = img.size
        if w > MAX_IMAGE_DIMENSION or h > MAX_IMAGE_DIMENSION:
            raise ValueError(
                f"Image dimensions {w}x{h} exceed max {MAX_IMAGE_DIMENSION}"
            )
        max_side = 2048
        if w > max_side or h > max_side:
            ratio = min(max_side / w, max_side / h)
            nw, nh = int(w * ratio), int(h * ratio)
            img = img.resize((nw, nh), Image.Resampling.LANCZOS)
            w, h = nw, nh
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode("ascii")
        media = "image/jpeg"
        thumb = self.generate_thumbnail(file_bytes)
        return ProcessedFile(
            type="image",
            base64_data=b64,
            media_type=media,
            dimensions={"width": w, "height": h},
            thumbnail_bytes=thumb,
        )

    def generate_thumbnail(
        self, file_bytes: bytes, max_dim: int = 400, media_type: Optional[str] = None
    ) -> bytes:
        """Generate JPEG thumbnail, max dimension 400px. Supports images and PDF first page."""
        if media_type == "application/pdf":
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                page = doc.load_page(0)
                mat = fitz.Matrix(2, 2)
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                doc.close()
            except (OSError, ValueError, TypeError) as e:
                raise ValueError("Could not render PDF for thumbnail") from e
        else:
            try:
                img = Image.open(io.BytesIO(file_bytes))
                if img.mode in ("RGBA", "P"):
                    img = img.convert("RGB")
                elif img.mode != "RGB":
                    img = img.convert("RGB")
            except (OSError, ValueError, TypeError) as e:
                raise ValueError("Could not open image for thumbnail") from e
        w, h = img.size
        if w > max_dim or h > max_dim:
            ratio = min(max_dim / w, max_dim / h)
            nw, nh = int(w * ratio), int(h * ratio)
            img = img.resize((nw, nh), Image.Resampling.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=75)
        return buf.getvalue()

    def compute_hash(self, file_bytes: bytes) -> str:
        """SHA-256 hex digest."""
        return hashlib.sha256(file_bytes).hexdigest()
